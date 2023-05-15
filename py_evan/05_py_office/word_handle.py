# coding:utf-8
# time: 2023/5/13
# author: evan
import glob
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt


def read_word_paragraph():
    doc = Document('./doc/文本.docx')
    # for p in doc.paragraphs:
    # print(p.text)
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                print(cell.text, end=' ')
            print()


class ReadDoc:
    def __init__(self, path):
        self.doc = Document(path)
        self.p_text = ''
        self.table_text = ''
        self.get_paragraph()
        self.get_table()

    def get_paragraph(self):
        for p in self.doc.paragraphs:
            self.p_text += p.text + '\n'

    def get_table(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.table_text += cell.text
                self.table_text += '\n'


def cv_filter(path, target):
    """
    简历筛选
    :return:
    """
    res = []

    result = glob.glob(path)
    for i in result:
        if os.path.isfile(i):
            if i.endswith(".docx"):
                doc = ReadDoc(i)
                p_text = doc.p_text
                p_table = doc.table_text

                for t in target:
                    if t not in p_text and t not in p_table:
                        break
                else:
                    res.append(i)

    return res


class WriteDoc:
    def __init__(self, path):
        self.path = path
        self.doc = Document(path)

    def add_header(self):
        title = self.doc.add_heading('My title hehehe', 1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style.font.size = Pt(80)
        _t = title.add_run('\n leon')
        _t.italic = True
        self.doc.save(self.path)

    def add_paragraph(self, content='hello'):
        self.doc.add_paragraph(content)
        self.doc.save(self.path)

    def add_picture(self, img_path):
        self.doc.add_picture(img_path, width=Inches(5), height=Inches(5))
        self.doc.save(self.path)

    def add_table(self, rows=1, cols=3):
        table = self.doc.add_table(rows, cols, style='Colorful Shading Accent 6')
        cells = table.rows[0].cells
        cells[0].text = 'a'
        cells[1].text = 'b'
        cells[2].text = 'c'
        self.doc.save(self.path)

    def add_page(self):
        self.doc.add_page_break()
        self.doc.save(self.path)

    @property
    def document(self):
        return self.doc


class DocStyleSetter:
    def __init__(self, doc: Document, path):
        self.doc = doc
        self.path = path
        self.style = self.doc.styles['Normal']

    def set_font(self):
        self.style.font.name = '微软雅黑'
        self.style.font.color.rgb = RGBColor(188, 44, 44)
        self.style.font.size = Pt(16)
        self.doc.save(self.path)

    def set_image(self, image_path):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _p = p.add_run()
        _p.add_run(image_path)


if __name__ == '__main__':
    # read_word_paragraph()
    '''
        path = os.path.join(os.getcwd(), 'doc', '*')
        l = cv_filter(path, ['python'])
        print(l)
        print(len(l))
    '''

    doc_writer = WriteDoc('./test_files/leon.docx')
    # doc_writer.add_header()
    # doc_writer.add_paragraph()
    # doc_writer.add_picture('./doc/logo2020.png')
    # doc_writer.add_table()
    # doc_writer.add_page()
    doc_writer.add_header()
    document = doc_writer.document

    setter = DocStyleSetter(document, './test_files/leon.docx')
    setter.set_font()
